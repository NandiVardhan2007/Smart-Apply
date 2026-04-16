import os
from docx import Document

def tag_template(input_path, output_path):
    doc = Document(input_path)
    
    # 1. Direct Text Replacements
    replacements = {
        "PEYYALA MANIDEEP": "{{ name }}",
        "vloggingwithmani@gmail.com": "{{ contact.email }}",
        "+911234567890": "{{ contact.phone }}",
        "NELLORE India": "{{ contact.location }}"
    }
    
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)
                
    # 2. Table Replacements
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            para.text = para.text.replace(old, new)

    # 3. Simple Summary Replacement (Look for the summary text and inject tag)
    summary_trigger = "Detail-oriented Data Analyst"
    for para in doc.paragraphs:
        if summary_trigger in para.text:
            para.text = "{{ summary }}"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if summary_trigger in para.text:
                        para.text = "{{ summary }}"

    # Save
    doc.save(output_path)
    print(f"Tagged template saved to {output_path}")

base_dir = r"D:\SMARTAPPLY\43_resume_templates"
dest_dir = r"d:\SMARTAPPLY\Backend\app\static\templates\resumes"

mapping = {
    "reusme_template_41.docx": "executive_gold.docx",
    "reusme_template_7.docx": "modern_premium.docx",
    "reusme_template_23.docx": "creative_premium.docx",
    "reusme_template_10.docx": "structured_standard.docx",
    "reusme_template_43.docx": "minimalist_sleek.docx"
}

for src, dst in mapping.items():
    tag_template(os.path.join(base_dir, src), os.path.join(dest_dir, dst))
