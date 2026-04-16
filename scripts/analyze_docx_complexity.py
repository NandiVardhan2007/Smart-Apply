import os
from docx import Document

base_dir = r"D:\SMARTAPPLY\43_resume_templates"
files = [f for f in os.listdir(base_dir) if f.endswith('.docx')]
files.sort()

results = []
for f in files:
    path = os.path.join(base_dir, f)
    try:
        doc = Document(path)
        stats = {
            "filename": f,
            "tables": len(doc.tables),
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections),
            "size_kb": round(os.path.getsize(path) / 1024, 2)
        }
        results.append(stats)
    except Exception as e:
        pass

# Sort by complexity (tables then size)
results.sort(key=lambda x: (x['tables'], x['size_kb']), reverse=True)

print("Top 10 Most Complex Templates (Potential Premium/Important):")
for r in results[:10]:
    print(r)

print("\nRepresentative Simple Templates:")
results.sort(key=lambda x: (x['tables'], x['size_kb']))
for r in results[:5]:
    print(r)
